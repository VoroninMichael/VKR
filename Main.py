import tkinter as tk
from tkinter import ttk,messagebox
from tkcalendar import Calendar,DateEntry
from datetime import date
import docx
from Reception import Shoes
from trash import Trash
from insert_db import Db
from Shoes_all import Shoes_all


from Task import Task

class Main(tk.Toplevel):
    def __init__(self,window,app,id_user):
        super().__init__(window)
        self.db=db
        self.id_user=id_user
        self.window=window
        self.app=app
        self.find_id_shop(self.id_user)
        self.init_main()



    def find_id_shop(self,id_employ):
        self.db.cursor.execute('''SELECT id_shop FROM employees where id=(%s) ''',(id_employ,))
        self.id_shop= int(str(self.db.cursor.fetchone()[0]))
    def init_main(self):
        self.title("Merchendaiser")
        self.geometry('650x350')

        self.resizable(False,False)
        Btn_empty=tk.Button(self,text='Отобразить пустые полки', command=self.find_empty)
        Btn_empty.place(x=10,y=110)
        Btn_Shoes=tk.Button(self,text='Склад',width=25, command=self.open_dialog_shoes)
        Btn_Shoes.place(x=10,y=150)
        Btn_Shoesall=tk.Button(self,text='Товар всех магазинов',width=25, command= self.shoes_all)
        Btn_Shoesall.place(x=10,y=180)
        Btn_write_off=tk.Button(self,text='Списанный/купленный товар', width=25,command= self.write_off_all)
        Btn_write_off.place(x=10,y=210)

        Btn_task=tk.Button(self,text='Создать задачу',width=25,command=self.task)
        Btn_task.place(x=10,y=240)
        Btn_plan=tk.Button(self,text='Отобразить план',width=25,command=self.create_plan)
        Btn_plan.place(x=10,y=270)


        Btn_sold_shell=tk.Button(self,text='Купить товар', command=self.sold)
        Btn_sold_shell.place(x=227,y=310)
        Btn_write_off_shell=tk.Button(self,text='Снять товар с полки', command=self.open_dialog_writeOff_shell)
        Btn_write_off_shell.place(x=317,y=310)



        Btn_close=tk.Button(self,text='Выход', command=self.out)
        Btn_close.place(x=3,y=310)

        self.zag1 = tk.Label(self, text="Выберите витрину", font=("Arial Bold", 10))
        self.zag1.place(x=1, y=1)

        self.combo1 = ttk.Combobox(self,width=27,state="readonly")

        self.combo1['values']=self.view_type()
        self.combo1.place(x=10, y=20)
        self.combo1.bind("<<ComboboxSelected>>",self.link)
        self.combo2 = ttk.Combobox(self,width=27,state="readonly")
        self.combo2['values']
        self.combo2.place(x=10, y=50)
        self.combo2.bind("<<ComboboxSelected>>",self.link_showcase)
        self.combo3 = ttk.Combobox(self,width=27,state="readonly")
        self.combo3['values']
        self.combo3.place(x=10, y=80)
        self.combo3.bind("<<ComboboxSelected>>",self.view_tree)


        self.tree= ttk.Treeview(self,columns=('ID','name','color','material','season','price','discount'),height=14,  show='headings')
        #

        self.tree.column('ID', width=20, anchor=tk.CENTER)
        self.tree.column('name', width=90, anchor=tk.CENTER)
        self.tree.column('color', width=60, anchor=tk.CENTER)
        self.tree.column('material', width=70, anchor=tk.CENTER)
        self.tree.column('season', width=70, anchor=tk.CENTER)
        self.tree.column('price', width=40, anchor=tk.CENTER)
        self.tree.column('discount', width=50, anchor=tk.CENTER)


        self.tree.heading('ID',text='id')
        self.tree.heading('name',text='Название')
        self.tree.heading('color',text='цвет')
        self.tree.heading('material',text='Матерьял')
        self.tree.heading('season',text='Сезон')
        self.tree.heading('price',text='Цена')
        self.tree.heading('discount',text='Скидка')
        self.ysb = ttk.Scrollbar(self, command=self.tree.yview)
        self.tree.configure(yscroll=self.ysb.set)
        self.ysb.pack(side=tk.RIGHT,fill="y")
        self.tree.place(x=227,y=3)


    def clear(self):
        self.tree.column('ID', width=20, anchor=tk.CENTER)
        self.tree.column('name', width=90, anchor=tk.CENTER)
        self.tree.column('color', width=60, anchor=tk.CENTER)
        self.tree.column('material', width=70, anchor=tk.CENTER)
        self.tree.column('season', width=70, anchor=tk.CENTER)
        self.tree.column('price', width=40, anchor=tk.CENTER)
        self.tree.column('discount', width=50, anchor=tk.CENTER)
        for i in self.tree.get_children():
            self.tree.delete(i)
        self.combo1.set("")
        self.combo2.set("")
        self.combo3.set("")


    def view_tree (self,select):
        self.view()

    def find_empty(self):

        self.tree.column('ID', width=150, anchor=tk.CENTER)
        self.tree.column('name', width=450, anchor=tk.CENTER)
        for i in self.tree.get_children():
            self.tree.delete(i)

        self.db.cursor.execute('''SELECT id FROM departament where id_shop = (%s)''', (self.id_shop,))
        for row in self.db.cursor.fetchall():
            self.db.cursor.execute('''SELECT name FROM departament where id = (%s)''', (row[0],))
            self.tree.insert('', tk.END, values=(("Отдел"),(self.db.cursor.fetchone()[0])))
            self.db.cursor.execute('''select depttype.id  from depttype
inner join departament ON depttype.id  = departament.id_depttype
where departament.id = (%s)''', (row[0],))
            for i in self.db.cursor.fetchall():
                self.db.cursor.execute('''SELECT name FROM depttype where id = (%s)''', (i[0],))
                self.tree.insert('', tk.END, values=(("Подотдел"),(self.db.cursor.fetchone()[0])))
                self.db.cursor.execute('''select subdept_showcase.id  from depttype
    inner join subdept_showcase ON depttype.id_subdept  = subdept_showcase.id
    where depttype.id = (%s)''', (i[0],))
                for j in self.db.cursor.fetchall():
                    self.db.cursor.execute('''SELECT name FROM subdept_showcase where id = (%s)''', (j[0],))
                    self.name_subdept_showcase=str(self.db.cursor.fetchone()[0])
                    self.tree.insert('', tk.END, values=((("Ветрина"),(self.name_subdept_showcase))))
                    self.db.cursor.execute('''SELECT id_showcase FROM subdept_showcase where name = (%s)''', (self.name_subdept_showcase,))
                    self.id_showcase=int(str(self.db.cursor.fetchone()[0]))
                    self.db.cursor.execute('''SELECT levelmax FROM showcase where id = (%s)''', (self.id_showcase,))
                    self.levelmax=int(str(self.db.cursor.fetchone()[0]))
                    for a in range(self.levelmax):
                        self.data=[]
                        self.db.cursor.execute('''SELECT id_shelf FROM showcase_shelf where id_showcase = (%s)''', (self.id_showcase,))
                        for wor in reversed(self.db.cursor.fetchall()):
                            self.data.append(wor[0])
                        self.db.cursor.execute('''SELECT count(*)  from shoes_shelf
                inner join shoes  ON shoes_shelf.id_shoes  = shoes.id
                inner join color on shoes.id=color.id_shoes
                where shoes_shelf.id_shelf= (%s) and shoes_shelf.id_shoes is not null ''', (self.data[a],))
                        self.occupied=int(str(self.db.cursor.fetchone()[0]))
                        if self.occupied >= 0:
                            self.db.cursor.execute('''SELECT kappa FROM shelf where id = (%s)''', (self.data[a],))
                            self.kappa=int(str(self.db.cursor.fetchone()[0])) - self.occupied
                            if self.kappa != 0 :
                                self.tree.insert('', tk.END, values=(a+1,"Полка:"))
                                for j in range(self.kappa):
                                    self.tree.insert('', tk.END, values=("","Пусто"))
                            if self.kappa == 0 :
                                self.tree.insert('', tk.END, values=("","Полка заполнена"))






    def view(self):
        self.db.cursor.execute('''SELECT id_showcase FROM subdept_showcase where name = (%s)''', (self.combo3.get(),))
        self.id_showcase=int(str(self.db.cursor.fetchone()[0]))
        self.db.cursor.execute('''SELECT levelmax FROM showcase where id = (%s)''', (self.id_showcase,))
        self.levelmax=int(str(self.db.cursor.fetchone()[0]))
        for i in self.tree.get_children():
            self.tree.delete(i)
        for i in range(self.levelmax):
            self.tree.insert('', tk.END, values=(i+1," Полка"))
            self.data=[]
            self.db.cursor.execute('''SELECT id_shelf FROM showcase_shelf where id_showcase = (%s)''', (self.id_showcase,))
            for row in reversed(self.db.cursor.fetchall()):
                self.data.append(row[0])
            self.db.cursor.execute('''SELECT count(*)  from shoes_shelf
inner join shoes  ON shoes_shelf.id_shoes  = shoes.id
inner join color on shoes.id=color.id_shoes where shoes_shelf.id_shelf= (%s) and shoes_shelf.id_shoes is not null ''', (self.data[i],))
            self.occupied=int(str(self.db.cursor.fetchone()[0]))
            self.db.cursor.execute('''SELECT id_shoes FROM shoes_shelf where id_shelf = (%s)''', (self.data[i],))
            for j in range(self.occupied):
                self.db.cursor.execute('''select shoes.id,shoes."name",color.color,shoes.material,shoes.season,shoes.price,shoes.discount  from shoes_shelf
inner join shoes  ON shoes_shelf.id_shoes  = shoes.id
inner join color on shoes.id=color.id_shoes where shoes_shelf.id_shelf = (%s)
 ''', (self.data[i],))
                self.tree.insert('', 'end', values=(self.db.cursor.fetchall()[j]))
            self.db.cursor.execute('''SELECT kappa FROM shelf where id = (%s)''', (self.data[i],))
            self.kappa=int(str(self.db.cursor.fetchone()[0])) - self.occupied
            for j in range(self.kappa):
                self.tree.insert('', tk.END, values=("","пусто"))
        self.tree.column('ID', width=20, anchor=tk.CENTER)
        self.tree.column('name', width=90, anchor=tk.CENTER)
        self.tree.column('color', width=60, anchor=tk.CENTER)
        self.tree.column('material', width=70, anchor=tk.CENTER)
        self.tree.column('season', width=70, anchor=tk.CENTER)
        self.tree.column('price', width=40, anchor=tk.CENTER)
        self.tree.column('discount', width=50, anchor=tk.CENTER)


    def link(self,select):

        self.name=str(self.combo1.get())
        self.db.cursor.execute('''SELECT id FROM departament where name = (%s) ''', (self.name,))

        self.id_departament = int(str(self.db.cursor.fetchone()[0]))

        self.combo2.set("")
        self.combo3.set("")
        self.db.cursor.execute('''select depttype."name" from departament
        inner join depttype  on depttype.id = departament.id_depttype where departament.id_shop=(%s) and departament.name=(%s) ''', (self.id_shop,self.name,))
        self.data=[]

        for row in self.db.cursor.fetchall():
            if row[0] not in self.data:
                self.data.append(row[0])
        self.combo2['values']=self.data
        return self.data

    def link_showcase(self,select):
        self.combo3.set("")
        self.db.cursor.execute('''SELECT id FROM depttype where name = (%s) ''', (self.combo2.get(),))
        self.id_subdepartament = int(str(self.db.cursor.fetchone()[0]))
        self.db.cursor.execute('''select subdept_showcase.name from depttype
inner join subdept_showcase  ON depttype.id_subdept  = subdept_showcase.id
inner join departament ON depttype.id=departament.id_depttype where depttype.id=(%s)  ''', (self.id_subdepartament,))
        self.data=[]
        for row in self.db.cursor.fetchall():
            if row[0] not in self.data:
                self.data.append(row[0])
        self.combo3['values']=self.data
        return self.data

    def open_dialog_shoes(self):
        self.clear()
        Shoes(self.window,self.id_shop)
    def shoes_all(self):
        self.clear()
        Shoes_all(self.window)
    def write_off_all(self):
        self.clear()
        Trash(self.window,self.id_shop)

    def task(self):
        self.clear()
        Task(self.window,self.id_shop)



    def view_type(self):
        self.db.cursor.execute('''SELECT name FROM departament where id_shop = (%s)''',(self.id_shop,))
        self.data=[]
        for row in self.db.cursor.fetchall():
            if row[0] not in self.data:
                self.data.append(row[0])
        return self.data

    def open_dialog_writeOff_shell(self):
        try:
            self.db.cursor.execute('''SELECT id_showcase FROM subdept_showcase where name = (%s)''', (self.combo3.get(),))
            self.id_showcase=int(str(self.db.cursor.fetchone()[0]))
            self.write_off=self.tree.set(self.tree.selection()[0], '#1')
            self.db.cursor.execute(''' select shoes_shelf.id from shoes_shelf
    inner join showcase_shelf on shoes_shelf.id_shelf = showcase_shelf.id_shelf
    inner join showcase on showcase.id  = showcase_shelf.id_showcase
    inner join shoes  ON shoes_shelf.id_shoes  = shoes.id
    inner join color on shoes.id=color.id_shoes where shoes_shelf.id_shoes=(%s) and showcase.id=(%s)   ''',(self.write_off, self.id_showcase,))
            self.id_shelf=self.id_shelf=int(str(self.db.cursor.fetchone()[0]))
            self.db.cursor.execute('''update shoes_shelf set id_shoes = default where id = (%s) ''', (self.id_shelf,))
            self.db.cursor.execute('''update shoes set count_shelf = count_shelf-1 ,count = count+1 where id = (%s) ''', (self.write_off,))
            messagebox.showinfo('Очень важно', 'Товар был снят с полки')
            self.view()
        except:
            messagebox.showinfo('Ошибка', 'Выберите товар, который хотите снять')

    def sold(self):
        try:
            self.db.cursor.execute('''SELECT id_showcase FROM subdept_showcase where name = (%s)''', (self.combo3.get(),))
            self.id_showcase=int(str(self.db.cursor.fetchone()[0]))
            self.buy=self.tree.set(self.tree.selection()[0], '#1')
            self.db.cursor.execute(''' select shoes_shelf.id from shoes_shelf
            inner join showcase_shelf on shoes_shelf.id_shelf = showcase_shelf.id_shelf
            inner join showcase on showcase.id  = showcase_shelf.id_showcase
            inner join shoes  ON shoes_shelf.id_shoes  = shoes.id
            inner join color on shoes.id=color.id_shoes where shoes_shelf.id_shoes=(%s) and showcase.id=(%s)   ''',(self.buy, self.id_showcase,))
            self.id_shelf=self.id_shelf=int(str(self.db.cursor.fetchone()[0]))
            self.db.cursor.execute('''select id_shoes from sold where id_shoes=(%s) ''',(self.buy,))
            if self.db.cursor.fetchone() is None :
                self.db.cursor.execute('''insert into sold (id_shoes,sold_date) values (%s,%s) ''',(self.buy,date.today(),))
            self.db.cursor.execute('''update shoes_shelf set id_shoes = default where id = (%s) ''', (self.id_shelf,))
            self.db.cursor.execute('''update shoes set count_shelf = count_shelf-1, sold_count=sold_count+1 where id = (%s) ''', (self.buy,))
            messagebox.showinfo('Очень важно', 'Товар был куплен')
            self.view()
        except:
            messagebox.showinfo('Ошибка', 'Выберите товар, который хотите купить')


    def create_plan(self):
        mydoc = docx.Document()
        self.db.cursor.execute('''SELECT address FROM shop where id = (%s)''', (self.id_shop,))
        mydoc.add_paragraph(("Магазин по адресу:",self.db.cursor.fetchone()[0]))
        self.db.cursor.execute('''SELECT id FROM departament where id_shop = (%s)''', (self.id_shop,))
        for row in self.db.cursor.fetchall():
            self.db.cursor.execute('''SELECT name FROM departament where id = (%s)''', (row[0],))
            mydoc.add_paragraph((("Отдел"),(self.db.cursor.fetchone()[0])))
            self.db.cursor.execute('''select depttype.id  from depttype
inner join departament ON depttype.id  = departament.id_depttype
where departament.id = (%s)''', (row[0],))
            for i in self.db.cursor.fetchall():
                self.db.cursor.execute('''SELECT name FROM depttype where id = (%s)''', (i[0],))
                mydoc.add_paragraph((("Подотдел"),(self.db.cursor.fetchone()[0])))
                self.db.cursor.execute('''select subdept_showcase.id  from depttype
    inner join subdept_showcase ON depttype.id_subdept  = subdept_showcase.id
    where depttype.id = (%s)''', (i[0],))
                for j in self.db.cursor.fetchall():
                    self.db.cursor.execute('''SELECT name FROM subdept_showcase where id = (%s)''', (j[0],))
                    self.name_subdept_showcase=str(self.db.cursor.fetchone()[0])
                    mydoc.add_paragraph((("Ветрина"),(self.name_subdept_showcase)))
                    self.db.cursor.execute('''SELECT id_showcase FROM subdept_showcase where name = (%s)''', (self.name_subdept_showcase,))
                    self.id_showcase=int(str(self.db.cursor.fetchone()[0]))
                    self.db.cursor.execute('''SELECT levelmax FROM showcase where id = (%s)''', (self.id_showcase,))
                    self.levelmax=int(str(self.db.cursor.fetchone()[0]))
                    for a in range(self.levelmax):
                        x=str(a+1)
                        mydoc.add_paragraph((x," Полка"))
                        self.data=[]
                        self.db.cursor.execute('''SELECT id_shelf FROM showcase_shelf where id_showcase = (%s)''', (self.id_showcase,))
                        for wor in reversed(self.db.cursor.fetchall()):
                            self.data.append(wor[0])
                        self.db.cursor.execute('''SELECT count(*)  from shoes_shelf
            inner join shoes  ON shoes_shelf.id_shoes  = shoes.id
            inner join color on shoes.id=color.id_shoes where shoes_shelf.id_shelf= (%s) and shoes_shelf.id_shoes is not null ''', (self.data[a],))
                        self.occupied=int(str(self.db.cursor.fetchone()[0]))
                        self.db.cursor.execute('''SELECT id_shoes FROM shoes_shelf where id_shelf = (%s)''', (self.data[a],))
                        for l in range(self.occupied):
                            self.db.cursor.execute('''select shoes.id,shoes."name",color.color,shoes.material,shoes.count  from shoes_shelf
            inner join shoes  ON shoes_shelf.id_shoes  = shoes.id
            inner join color on shoes.id=color.id_shoes where shoes_shelf.id_shelf = (%s)
             ''', (self.data[a],))
                            mydoc.add_paragraph(str(self.db.cursor.fetchall()[l]))
                        self.db.cursor.execute('''SELECT kappa FROM shelf where id = (%s)''', (self.data[a],))
                        self.kappa=int(str(self.db.cursor.fetchone()[0])) - self.occupied
                        for l in range(self.kappa):
                            mydoc.add_paragraph(" пусто")
            mydoc.save('Магазин.docx')



    def out(self):
        self.clear()
        self.app.invizeoff()
        self.destroy()


db=Db()
